from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[1]
WORKSPACE_ROOT = SCRIPT_PATH.parents[3]
REPORT_DIR = PROJECT_ROOT / "report"
REPORT_ASSETS = REPORT_DIR / "assets"

BODY_FONT = "宋体"
LATIN_FONT = "Times New Roman"
CODE_FONT = "Consolas"


def font_path(bold: bool = False) -> str:
    candidate = Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc")
    if candidate.exists():
        return str(candidate)
    return str(Path("C:/Windows/Fonts/simhei.ttf"))


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold), size)


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                  font: ImageFont.FreeTypeFont, fill: str = "#344C5A") -> None:
    left, top, right, bottom = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=6)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text,
                        font=font, fill=fill, align="center", spacing=6)


def box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str,
        fill: str, outline: str = "#344C5A", size: int = 34, bold: bool = False) -> None:
    draw.rounded_rectangle(rect, radius=24, fill=fill, outline=outline, width=4)
    centered_text(draw, rect, text, load_font(size, bold), outline)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
          fill: str = "#289BB0", width: int = 6) -> None:
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = (x2, y2)
    base = (x2 - ux * 22, y2 - uy * 22)
    points = [tip, (base[0] + px * 12, base[1] + py * 12),
              (base[0] - px * 12, base[1] - py * 12)]
    draw.polygon(points, fill=fill)


def create_module_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1120), "#FFF8ED")
    draw = ImageDraw.Draw(image)
    centered_text(draw, (0, 25, 1800, 100), "《萌泡大作战》功能模块图", load_font(50, True))
    box(draw, (610, 120, 1190, 230), "萌泡大作战（单人版）", "#66D5E8", size=40, bold=True)
    groups = [
        ("场景与界面", "主菜单\n角色选择\n说明/暂停\n结算与退出", "#DDF5F9"),
        ("游戏核心", "玩家移动\n碰撞检测\n水泡与水浪\n木箱与道具", "#FFE8D9"),
        ("敌人与关卡", "随机地图\n巡逻敌人\n追踪敌人\n多血量 Boss", "#E9DDF6"),
        ("资源与反馈", "角色/道具图集\n背景音乐\n战斗音效\n粒子动画", "#E2F4E9"),
        ("数据与测试", "生命/分数\n关卡统计\nSTL 容器\n自动截图验收", "#FFF0C7"),
    ]
    width = 300
    gap = 35
    left0 = 80
    for index, (title, detail, color) in enumerate(groups):
        left = left0 + index * (width + gap)
        center = left + width // 2
        arrow(draw, (900, 230), (center, 350))
        box(draw, (left, 350, left + width, 455), title, color, size=33, bold=True)
        box(draw, (left, 495, left + width, 970), detail, "#FFFDF8", size=31)
    centered_text(draw, (100, 1000, 1700, 1080),
                  "输入、对象更新、碰撞、绘制和音频由 MoeBubbleGame 统一调度",
                  load_font(30), "#5C717C")
    image.save(path)


def create_flow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1500, 1900), "#FFF8ED")
    draw = ImageDraw.Draw(image)
    centered_text(draw, (0, 20, 1500, 100), "游戏主流程图", load_font(50, True))
    nodes = [
        ((500, 120, 1000, 215), "启动并加载资源", "#DDF5F9"),
        ((500, 285, 1000, 380), "主菜单", "#66D5E8"),
        ((500, 450, 1000, 545), "角色选择 / 操作说明", "#FFF0C7"),
        ((500, 615, 1000, 710), "初始化关卡与对象容器", "#E2F4E9"),
        ((500, 780, 1000, 875), "输入 → 更新 → 碰撞 → 清理 → 绘制", "#FFE8D9"),
        ((500, 965, 1000, 1060), "玩家生命是否归零？", "#FADCE2"),
        ((500, 1160, 1000, 1255), "本关敌人是否清空？", "#E9DDF6"),
        ((500, 1355, 1000, 1450), "是否已经完成第 3 关？", "#E9DDF6"),
        ((170, 1580, 610, 1680), "失败结算 / 重新挑战", "#FFBDB5"),
        ((890, 1580, 1330, 1680), "最终通关 / 返回菜单", "#BFEAD9"),
    ]
    for rect, text, color in nodes:
        box(draw, rect, text, color, size=31, bold="？" not in text)
    for i in range(7):
        first = nodes[i][0]
        second = nodes[i + 1][0]
        arrow(draw, ((first[0] + first[2]) // 2, first[3]),
              ((second[0] + second[2]) // 2, second[1]))
    arrow(draw, (500, 1012), (390, 1012))
    draw.line([(390, 1012), (390, 1540)], fill="#E75D72", width=6)
    arrow(draw, (390, 1540), (390, 1580), fill="#E75D72")
    centered_text(draw, (300, 950, 480, 1005), "是", load_font(28, True), "#E75D72")
    centered_text(draw, (1010, 1065, 1110, 1120), "否", load_font(28, True), "#5C717C")
    draw.line([(1000, 1208), (1190, 1208), (1190, 830), (1000, 830)], fill="#289BB0", width=6)
    arrow(draw, (1190, 830), (1000, 830))
    centered_text(draw, (1030, 1160, 1170, 1210), "否", load_font(28, True), "#289BB0")
    draw.line([(500, 1402), (330, 1402), (330, 1208), (500, 1208)], fill="#289BB0", width=6)
    arrow(draw, (330, 1208), (500, 1208))
    centered_text(draw, (340, 1345, 470, 1395), "否：下一关", load_font(26, True), "#289BB0")
    arrow(draw, (750, 1450), (1110, 1580), fill="#54A98A")
    centered_text(draw, (780, 1455, 1030, 1520), "是", load_font(28, True), "#54A98A")
    draw.line([(390, 1680), (390, 1790), (80, 1790), (80, 332)], fill="#A88BC7", width=5)
    arrow(draw, (80, 332), (500, 332), fill="#A88BC7")
    draw.line([(1110, 1680), (1110, 1790), (1420, 1790), (1420, 332)], fill="#A88BC7", width=5)
    arrow(draw, (1420, 332), (1000, 332), fill="#A88BC7")
    centered_text(draw, (560, 1795, 940, 1860), "重新开始或返回主菜单", load_font(27), "#5C717C")
    image.save(path)


def create_class_diagram(path: Path) -> None:
    image = Image.new("RGB", (1900, 1250), "#FFF8ED")
    draw = ImageDraw.Draw(image)
    centered_text(draw, (0, 20, 1900, 95), "主要类关系与 STL 容器", load_font(48, True))
    box(draw, (720, 115, 1180, 220), "GameObject\n+ update()  + draw()  + bounds()", "#DDF5F9", size=23, bold=True)
    box(draw, (720, 305, 1180, 410), "Character\n位置 / 速度 / 碰撞 / move()", "#E2F4E9", size=27, bold=True)
    arrow(draw, (950, 305), (950, 220), fill="#344C5A")
    box(draw, (250, 505, 680, 620), "Player\n生命 / 角色属性 / 放泡能力", "#FFE8D9", size=27, bold=True)
    box(draw, (750, 505, 1240, 620), "Enemy（抽象类）\nchooseDirection() / takeWaveHit()", "#E9DDF6", size=22, bold=True)
    arrow(draw, (465, 505), (820, 410), fill="#344C5A")
    arrow(draw, (995, 505), (995, 410), fill="#344C5A")
    enemies = [
        ("PatrolEnemy\n随机巡逻与避险", (360, 760, 720, 865), "#E9DDF6"),
        ("HunterEnemy\n追踪玩家与避险", (770, 760, 1130, 865), "#FADCE2"),
        ("BossEnemy\n多血量 / 多态受击", (1180, 760, 1540, 865), "#FFF0C7"),
    ]
    for text, rect, color in enemies:
        box(draw, rect, text, color, size=25, bold=True)
        arrow(draw, ((rect[0] + rect[2]) // 2, rect[1]), (995, 620), fill="#344C5A")
    box(draw, (35, 120, 570, 410), "MoeBubbleGame（组合根）\n\nstd::vector<WaterBubble>\nstd::vector<WaterWave>\nstd::vector<PowerUp>\nstd::vector<StarParticle>\nstd::vector<std::unique_ptr<Enemy>>", "#FFFDF8", size=21, bold=True)
    arrow(draw, (570, 255), (720, 170), fill="#289BB0")
    box(draw, (1300, 120, 1855, 335), "固定与随机数据\n\nstd::array<TileType, 195>\nstd::array<bool, 256>\nstd::mt19937\nstd::uniform_*_distribution", "#FFFDF8", size=25, bold=True)
    arrow(draw, (1300, 230), (1180, 170), fill="#289BB0")
    box(draw, (110, 990, 1790, 1160),
        "运行时多态：主循环只持有 Enemy 基类智能指针；调用 draw、chooseDirection、takeWaveHit 时，\n"
        "由实际对象决定执行 PatrolEnemy、HunterEnemy 或 BossEnemy 的重写版本。unique_ptr 负责自动释放对象。",
        "#DDF5F9", size=27, bold=True)
    image.save(path)


def remove_all_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_east_asia_font(run, east_asia: str, latin: str | None = None) -> None:
    latin = latin or east_asia
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def format_body(paragraph: Paragraph, indent: bool = True) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(24) if indent else Pt(0)
    for run in paragraph.runs:
        set_east_asia_font(run, BODY_FONT, LATIN_FONT)
        run.font.size = Pt(12)


def set_body(paragraph: Paragraph, text: str, indent: bool = True) -> Paragraph:
    remove_all_runs(paragraph)
    paragraph.add_run(text)
    format_body(paragraph, indent)
    return paragraph


def new_paragraph_before(target: Paragraph, text: str = "", style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def add_body_before(target: Paragraph, text: str, indent: bool = True) -> Paragraph:
    paragraph = new_paragraph_before(target, text)
    format_body(paragraph, indent)
    return paragraph


def add_caption_before(target: Paragraph, text: str) -> Paragraph:
    paragraph = new_paragraph_before(target, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = False
    for run in paragraph.runs:
        set_east_asia_font(run, BODY_FONT, LATIN_FONT)
        run.font.size = Pt(12)
    return paragraph


def add_image_before(target: Paragraph, image_path: Path, width_inches: float = 5.55) -> Paragraph:
    paragraph = new_paragraph_before(target)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return paragraph


def add_code_before(target: Paragraph, code: str) -> None:
    for line in code.strip("\n").splitlines():
        paragraph = new_paragraph_before(target, line)
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.right_indent = Pt(12)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F3F5F6")
        paragraph._p.get_or_add_pPr().append(shading)
        for run in paragraph.runs:
            set_east_asia_font(run, "微软雅黑", CODE_FONT)
            run.font.size = Pt(9)


def paragraph_by_text(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise KeyError(f"paragraph not found: {text}")


def set_heading_text(paragraph: Paragraph, text: str) -> None:
    remove_all_runs(paragraph)
    run = paragraph.add_run(text)
    set_east_asia_font(run, "黑体", LATIN_FONT)
    run.bold = False


def set_cover_field(paragraph: Paragraph, label: str, value: str) -> None:
    remove_all_runs(paragraph)
    label_run = paragraph.add_run(label)
    set_east_asia_font(label_run, BODY_FONT, LATIN_FONT)
    label_run.font.size = Pt(16)
    label_run.bold = True
    value_run = paragraph.add_run("   " + value + "   ")
    set_east_asia_font(value_run, BODY_FONT, LATIN_FONT)
    value_run.font.size = Pt(16)
    value_run.bold = True
    value_run.underline = True


def set_cell_text(cell, text: str, center: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_stl_table(document: Document, before: Paragraph) -> None:
    caption = new_paragraph_before(before, "表2.1 STL数据结构与操作数据")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        set_east_asia_font(run, BODY_FONT, LATIN_FONT)
        run.font.size = Pt(12)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table._tbl.getparent().remove(table._tbl)
    before._p.addprevious(table._tbl)
    headers = ["STL组件", "保存/处理的数据", "设计理由"]
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, True, 10)
        table.rows[0].cells[index]._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        table.rows[0].cells[index]._tc.tcPr[-1].set(qn("w:fill"), "DDF5F9")
    set_repeat_table_header(table.rows[0])
    rows = [
        ("std::array", "195个地图格、256个键盘状态", "规模固定，连续存储并可按下标O(1)访问"),
        ("std::vector", "水泡、水浪、道具、粒子、危险格", "对象数量随游戏动态变化，便于追加、遍历和擦除"),
        ("vector<unique_ptr<Enemy>>", "巡逻、追踪、Boss异构敌人", "RAII自动释放，并通过基类指针实现运行时多态"),
        ("std::mt19937", "木箱、掉落、AI方向、粒子参数", "标准随机分布与固定种子便于复现实验"),
        ("std::algorithm", "查找、约束、选择和对象清理", "使用clamp/find/remove_if/min_element减少手写循环错误"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, index == 0, 9.5)
    widths = [Inches(1.35), Inches(2.05), Inches(2.25)]
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def fill_test_table(table) -> None:
    tests = [
        ("T01", "四配置编译", "生成Debug/Release、Win32/x64", "MSBuild", "四种配置均生成EXE并复制资源", "通过"),
        ("T02", "主菜单资源", "启动并执行报告截图模式", "无", "标题、原画、按钮及中文正常显示", "通过"),
        ("T03", "角色差异", "依次选择四名角色进入游戏", "鼠标/键盘", "生命、速度、容量、范围、护盾符合配置", "待人工复核"),
        ("T04", "鼠标按钮", "点击主菜单、角色、暂停、结算按钮", "鼠标", "悬停选中且只触发对应按钮", "待人工复核"),
        ("T05", "键盘兼容", "用方向键/WASD和Enter操作菜单", "键盘", "原键盘操作不受鼠标支持影响", "待人工复核"),
        ("T06", "地图边界", "持续向墙、木箱和边缘移动", "WASD", "脚底碰撞盒不穿墙且不离开地图", "待人工复核"),
        ("T07", "窄道转角", "沿墙提前输入垂直方向", "组合方向键", "自动向通道中心收拢并顺滑转弯", "待人工复核"),
        ("T08", "低摆幅步态", "直线移动并反复启停", "方向键", "回中帧衔接，人物无大幅左右摇摆", "待人工复核"),
        ("T09", "初始水泡容量", "未强化时连续放置水泡", "多次Space", "不超过角色初始容量，爆炸后归还", "待人工复核"),
        ("T10", "水浪边界", "在固定墙和木箱旁引爆", "Space", "墙完全阻挡；木箱破坏并终止传播", "待人工复核"),
        ("T11", "连锁触发", "同一直线上放置两枚水泡", "Space", "第一枚水浪立即触发第二枚且不重复", "待人工复核"),
        ("T12", "道具图标", "查看说明、地图掉落和HUD", "无", "四类透明图标清晰且无粉/黑底", "通过（截图）"),
        ("T13", "强化上限", "重复拾取同类道具", "角色碰撞", "范围≤6、容量≤5、速度≤4级、盾≤2", "待人工复核"),
        ("T14", "受伤与护盾", "接触水浪或敌人", "碰撞", "护盾优先抵消，随后扣生命并短暂无敌", "待人工复核"),
        ("T15", "普通敌人", "引导巡逻/追踪敌人进入水浪", "水浪", "敌人淘汰并增加300分", "待人工复核"),
        ("T16", "Boss血量", "进入第三关并查看顶部/头顶血条", "无", "最大生命10，两处血量一致", "通过（截图）"),
        ("T17", "Boss多段受击", "使用不同水泡连续命中Boss", "水浪", "一次有效水浪扣1点，归零加1500分", "待人工复核"),
        ("T18", "暂停与重开", "按P暂停并选择重开本关", "P/鼠标", "对象与音乐暂停，重开恢复本关初始状态", "待人工复核"),
        ("T19", "图片路径", "复制EXE及assets到新目录后启动", "文件系统", "全部PNG按EXE相对路径加载，无绝对路径依赖", "通过（构建）"),
        ("T20", "通关结算", "完成第三关或执行截图模式", "无", "显示分数、用时、木箱、道具和敌人统计", "通过（截图）"),
    ]
    while len(table.rows) < len(tests) + 1:
        table.add_row()
    while len(table.rows) > len(tests) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    headers = ["用例编号", "测试项", "测试步骤", "输入信号", "预期结果", "测试结果"]
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, True, 9)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(tests, start=1):
        for column_index, value in enumerate(values):
            set_cell_text(table.rows[row_index].cells[column_index], value,
                          column_index in (0, 3, 5), 8.5)


def count_han(texts: list[str]) -> int:
    return sum(1 for char in "".join(texts) if "\u4e00" <= char <= "\u9fff")


def build_report(args) -> Path:
    template = Path(args.template).resolve()
    output = Path(args.output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    REPORT_ASSETS.mkdir(parents=True, exist_ok=True)

    module_diagram = REPORT_ASSETS / "diagram_function_modules.png"
    flow_diagram = REPORT_ASSETS / "diagram_game_flow.png"
    class_diagram = REPORT_ASSETS / "diagram_class_stl.png"
    create_module_diagram(module_diagram)
    create_flow_diagram(flow_diagram)
    create_class_diagram(class_diagram)

    shutil.copy2(template, output)
    document = Document(output)

    # Remove only the placeholder/example images that came from the teacher's
    # template.  This must happen before inserting any report images; removing
    # every drawing later would also delete the generated diagrams and actual
    # program screenshots.
    template_image_relationships: set[str] = set()
    for paragraph in list(document.paragraphs):
        if paragraph._p.xpath(".//w:drawing"):
            for blip in paragraph._p.xpath(".//a:blip"):
                relationship_id = blip.get(qn("r:embed"))
                if relationship_id:
                    template_image_relationships.add(relationship_id)
            paragraph._element.getparent().remove(paragraph._element)
    for relationship_id in template_image_relationships:
        document.part.drop_rel(relationship_id)

    set_cover_field(document.paragraphs[11], "题    目", "萌泡大作战（单人闯关游戏）")
    set_cover_field(document.paragraphs[18], "序    号", args.sequence)
    set_cover_field(document.paragraphs[19], "学    号", args.student_id)
    set_cover_field(document.paragraphs[20], "姓    名", args.student_name)
    set_cover_field(document.paragraphs[21], "任课教师", args.teacher)
    set_cover_field(document.paragraphs[22], "成    绩", "                         ")
    set_body(document.paragraphs[30], "2026年7月", False)
    document.paragraphs[30].alignment = WD_ALIGN_PARAGRAPH.CENTER

    link = paragraph_by_text(document, "相关材料（电子报告及项目源码压缩包）的百度网盘链接：")
    set_body(link, "相关材料（电子报告及项目源码压缩包）的百度网盘链接：" + args.material_link, False)
    link.paragraph_format.space_after = Pt(12)

    set_heading_text(paragraph_by_text(document, "1.1 第一周学习总结（不少于1000汉字）"),
                     "1.1 第一周学习总结（除去代码，不少于1000汉字）")
    set_heading_text(paragraph_by_text(document, "1.2 第二周学习总结（同上）"),
                     "1.2 第二周学习总结（除去代码，不少于1000汉字）")

    week1_intro = [
        "第一周的学习重点是完成Visual Studio与EasyX开发环境的搭建，并建立图形程序与普通控制台程序之间的整体认识。EasyX把Windows图形窗口、绘图设备和消息机制封装成较容易上手的函数接口，通过initgraph创建固定尺寸窗口，通过closegraph释放图形资源。开发时需要正确选择Visual Studio的C++桌面开发组件和与平台一致的EasyX库，并理解Debug、Release、Win32和x64是不同的构建配置。项目采用解决方案统一管理多个练习，每个练习保持独立入口，既便于逐项调试，也避免多个main函数互相冲突。",
        "绘图部分首先掌握了物理坐标与逻辑坐标。EasyX默认以窗口左上角为原点，横向向右、纵向向下，所有图形最终都要转换成像素坐标。圆、椭圆、矩形、圆角矩形、多边形、直线和圆弧等基本图元虽然接口简单，但需要先确定窗口尺寸、边距、重复间隔和颜色，再通过循环计算每个元素的位置。设置线色、填充色和背景色时要区分setlinecolor、setfillcolor与setbkcolor的作用；绘制有边框与无边框图形时，也要根据需要选择fill系列或solid系列函数。五子棋棋盘就是把固定尺寸、边距、网格间距和星位坐标抽象成参数后重复绘制的典型案例。",
        "文字输出和动态刷新是第一周的另一个重点。settextstyle、settextcolor、textwidth和textheight共同决定字体、颜色与居中位置；使用TCHAR或宽字符串时必须让项目字符集和函数参数类型保持一致。对于会持续刷新的画面，应使用BeginBatchDraw和FlushBatchDraw，把一帧内容先绘制到缓冲区再整体显示，从而减少直接清屏造成的闪烁。动态时钟每帧读取系统时间，根据小时、分钟和秒数计算角度，再把角度转换为指针端点。这个过程把数学坐标、图形绘制、时间获取和对象职责结合起来，为后续游戏循环奠定了基础。",
    ]
    week1_difficult = [
        "第一周最容易出现的问题是字符集和图片路径。补充资料指出，loadimage编译报错通常是项目字符集与字符串实参不匹配；即使编译通过，图片不显示也常由格式或路径造成。我的处理方法是整个解决方案统一使用Unicode字符集，涉及路径时使用std::wstring和宽字符串接口；资源文件统一采用PNG/JPG等EasyX支持的位图格式。最终项目没有写死C盘或D盘绝对路径，而是先取得EXE所在目录，再拼接assets下的英文子目录，同时由项目文件在编译后自动复制资源。这种做法保证工程移动到老师电脑后仍能找到图片。",
        "第二个难点是把重复坐标从大量常数中提炼出来。最初直接写每条棋盘线或每个刻度的位置，不仅代码冗长，改变窗口尺寸时也很容易漏改。解决办法是把窗口宽高、边距、行数和单元尺寸作为类内常量，再用循环和统一公式计算坐标。动态时钟把角度减去90度后转换为弧度，使十二点方向与数学三角函数的零度方向对齐；小时、分钟和秒针分别封装长度、线宽和颜色，避免三个指针共享大量重复代码。",
        "第三个难点是画面刷新和资源生命周期。循环中若直接cleardevice再逐个绘制，肉眼会看到闪烁；若忘记处理退出消息，窗口会无法正常结束；若忘记closegraph，则资源释放不完整。通过固定一帧的顺序为“处理输入、更新数据、清除缓冲、绘制全部对象、刷新缓冲、限制帧率”，动态画面变得稳定。练习进一步使用类的构造函数建立有效初始状态，并把窗口初始化、绘制、输入等待和关闭操作封装在应用类中，使main函数只负责创建对象并调用run。",
    ]
    week1_cases = [
        "案例一是15×15五子棋棋盘。GobangBoard类把窗口尺寸、网格数量、边距和单元格尺寸封装为私有常量，对外只公开run和报告截图接口。drawGrid通过一次循环同时绘制横线与竖线，drawStarPoints根据3、7、11三个索引组合得到九个星位。该实现没有把每条线硬编码，调整cellSize即可整体改变棋盘比例，体现了数据封装与职责分解。",
        "案例二是动态模拟时钟。ClockHand类保存指针长度、线宽和颜色，并提供统一draw方法；DynamicClock组合三个ClockHand对象，读取SYSTEMTIME后分别计算时针、分针和秒针角度。表盘绘制使用60次循环区分小时刻度与分钟刻度，十二个数字利用textwidth和textheight居中。程序采用批量绘图并保持约60帧刷新，秒针能够平滑包含毫秒变化。",
    ]

    week2_intro = [
        "第二周从静态绘图进入交互式图形程序，重点学习EasyX消息结构、鼠标与键盘输入、持续动画、碰撞检测和面向对象协作。ExMessage能够描述鼠标按下、移动、键盘按下与抬起等事件；peekmessage适合在游戏循环中非阻塞地取出消息，getmessage则适合等待一次输入。对于菜单中的单次确认，需要识别从未按到按下的边沿；对于角色移动，需要同时保留按住状态和极短点击，避免只轮询某一种状态造成输入丢失。最终游戏的InputManager使用三个固定长度std::array记录当前帧、上一帧和两帧之间出现过的短按信号。",
        "动画程序应把逻辑更新与画面绘制分开。弹球练习以约16毫秒为目标帧间隔，每帧先处理输入，再更新球和挡板位置，最后绘制并刷新。球与墙碰撞时要把位置修正回合法边界，再改变速度方向，否则高速运动可能让球连续卡在墙内；球与挡板碰撞还要判断上一帧是否位于挡板上方，防止从挡板侧面或下方反复触发。根据球击中挡板中心的偏移比例调整水平速度，可以让操作产生可观察的控制效果，而不是每次按照固定角度反弹。",
        "面向对象设计在第二周由封装扩展到继承和多态。鼠标键盘绘图程序定义抽象Shape接口，Square和Circle分别重写draw；DrawingApp只接收Shape引用并调用draw，不需要为每一种图形重复编写消息流程。弹球程序则采用组合关系：BouncingBallGame拥有Ball与Paddle，由应用对象负责帧循环，Ball负责速度和碰撞，Paddle负责按键状态和边界。不同对象只暴露必要方法，内部坐标和状态保持私有，为最终项目中更完整的GameObject和Enemy继承体系提供了直接经验。",
    ]
    week2_difficult = [
        "第二周首先解决的是短按与长按的差异。只处理WM_KEYDOWN会让持续移动依赖系统键盘重复速度，只调用GetAsyncKeyState的高位又可能漏掉发生在两帧之间的快速点击。最终做法是把down用于连续移动，把pressed用于菜单和补偿短按，并在每帧保存previous状态。鼠标按钮使用窗口消息记录坐标和本帧左键按下标志，所有可见按钮共享矩形命中测试和悬停选中逻辑，因此键盘与鼠标可以同时工作而不会复制场景跳转代码。",
        "其次是连续碰撞与边界值。挡板必须限制在0到窗口宽度减挡板宽度之间，球的圆心必须考虑半径，游戏角色的碰撞盒则不能直接等同于帽子和背包的可见范围。最终项目使用约20×20像素的脚底碰撞盒，显示精灵仍可覆盖44像素步道；转角时根据角色与当前格中心的偏差做有限纠正，使碰撞严格性和操作流畅度兼顾。人物动画在左右跨步之间插入回中帧，并缩小跨步帧与校正重心，解决了原来左右摆动过大的问题。",
        "再次是动态对象存储和清理。水泡、水浪、道具和粒子数量随操作不断变化，固定数组难以管理，因此使用std::vector追加对象；对象失效后先标记active=false，再在统一清理阶段用std::remove_if和erase删除，避免遍历过程中直接删除导致迭代器失效。敌人包含多个派生类型，使用std::vector<std::unique_ptr<Enemy>>既能通过基类指针调用虚函数，又能在擦除时自动释放真实派生对象。随机地图、道具掉落和AI方向使用std::mt19937与标准分布，并设置固定种子，便于重复测试同一关卡。",
    ]
    week2_cases = [
        "案例一是鼠标键盘绘图程序。左键绘制正方形、右键绘制圆形，R/G/B/W切换颜色，按住Ctrl增大图形。Shape作为抽象基类只规定draw行为，Square和Circle重写具体图元；DrawingApp负责窗口与消息分发。新增图形类型时只需实现Shape，而不需要修改鼠标事件的主要流程，这一案例直观体现了接口抽象和运行时多态。",
        "案例二是键盘控制弹球。Paddle分别记录向左和向右的按住状态并用std::clamp限制边界；Ball封装坐标、速度、半径和碰撞方法；BouncingBallGame组合两个对象并维护暂停和帧率。球与挡板接触点决定新的水平速度，P键仅在首次按下时切换暂停，Esc安全退出。该案例将持续输入、时间步进、碰撞边界和类间协作集中在一个小型实时程序中。",
    ]

    markers = {
        "1.1.1 内容简介": week1_intro,
        "1.1.2 难点和解决办法": week1_difficult,
        "1.2.1 内容简介": week2_intro,
        "1.2.2 难点和解决办法": week2_difficult,
    }
    placeholder_for_heading = {
        "1.1.1 内容简介": "本周所学内容",
        "1.1.2 难点和解决办法": "学习中遇到的问题以及个人是如何解决的",
        "1.2.1 内容简介": "同上",
    }
    for heading_text, paragraphs in markers.items():
        heading = paragraph_by_text(document, heading_text)
        all_paragraphs = document.paragraphs
        index = next(i for i, p in enumerate(all_paragraphs) if p._p is heading._p)
        next_heading = next(p for p in all_paragraphs[index + 1:] if p.style.name.startswith("Heading"))
        body_candidates = [p for p in all_paragraphs[index + 1:] if p._p.getparent() is heading._p.getparent()]
        first = next(p for p in body_candidates if p._p is not next_heading._p)
        set_body(first, paragraphs[0])
        for text in paragraphs[1:]:
            add_body_before(next_heading, text)

    # The second occurrence of “同上” belongs to 1.2.2 and is filled explicitly.
    h122 = paragraph_by_text(document, "1.2.2 难点和解决办法")
    h123 = paragraph_by_text(document, "1.2.3 学习案例")
    between = [p for p in document.paragraphs if p._p.getparent() is h122._p.getparent()
               and h122._p.getparent().index(h122._p) < h122._p.getparent().index(p._p)
               < h122._p.getparent().index(h123._p)]
    set_body(between[0], week2_difficult[0])
    for text in week2_difficult[1:]:
        add_body_before(h123, text)

    h113 = paragraph_by_text(document, "1.1.3 学习案例")
    h12 = next(p for p in document.paragraphs if p.text.startswith("1.2 第二周"))
    case_placeholder = next(p for p in document.paragraphs
                            if h113._p.getparent().index(h113._p) < h113._p.getparent().index(p._p)
                            < h113._p.getparent().index(h12._p))
    set_body(case_placeholder, week1_cases[0])
    add_image_before(h12, REPORT_ASSETS / "exercise_basic_gobang.png", 4.7)
    add_caption_before(h12, "图1.1 EasyX基础练习：五子棋棋盘")
    add_body_before(h12, week1_cases[1])
    add_image_before(h12, REPORT_ASSETS / "exercise_basic_clock.png", 4.7)
    add_caption_before(h12, "图1.2 EasyX基础练习：动态模拟时钟")

    h123 = paragraph_by_text(document, "1.2.3 学习案例")
    h2 = paragraph_by_text(document, "2 项目开发报告")
    case2_placeholder = next(p for p in document.paragraphs
                             if h123._p.getparent().index(h123._p) < h123._p.getparent().index(p._p)
                             < h123._p.getparent().index(h2._p))
    set_body(case2_placeholder, week2_cases[0])
    add_image_before(h2, REPORT_ASSETS / "exercise_advanced_drawing.png", 5.55)
    add_caption_before(h2, "图1.3 EasyX进阶练习：鼠标键盘绘图")
    add_body_before(h2, week2_cases[1])
    add_image_before(h2, REPORT_ASSETS / "exercise_advanced_ball.png", 5.55)
    add_caption_before(h2, "图1.4 EasyX进阶练习：键盘控制弹球")

    project_intro = (
        "《萌泡大作战》是使用C++17、Visual Studio与EasyX开发的原创Q版单人网格闯关游戏。玩家扮演泡泡工坊搭档，在15×13格地图中放置净水泡泡；泡泡倒计时结束后向四个方向释放水浪，能够破坏木箱、触发其他泡泡并淘汰AI敌人。游戏包含三关、四名差异化角色、巡逻/追踪两类普通敌人和第三关多血量Boss“墨潮王”。四类角色分别偏向全能、速度、爆破与生存，四类道具可强化爆破范围、水泡容量、移动速度和护盾。项目提供主菜单、角色选择、操作说明、暂停、关卡结算、失败、通关与退出确认界面，同时包含原创角色图、行走精灵、道具图标、背景音乐和音效。故事设定为玩家净化被墨潮污染的云朵岛心泉，连续完成三处水域后恢复水循环。"
    )
    set_body(paragraph_by_text(document, "此部分对所选项目进行简要说明"), project_intro)

    module_intro = paragraph_by_text(document, "此部分要求绘制系统功能模块图，例如：")
    set_body(module_intro, "系统按职责划分为场景与界面、游戏核心、敌人与关卡、资源与反馈、数据与测试五个模块。MoeBubbleGame作为组合根维护主循环和场景状态，各模块通过公开接口协作，避免全局变量直接耦合。")
    caption21 = paragraph_by_text(document, "图2.1 功能模块图")
    set_body(caption21, "图2.1 游戏功能模块图", False)
    caption21.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption21, module_diagram, 5.55)

    requirements = [
        "场景需求：程序启动后进入主菜单，可以开始游戏、选择角色、查看说明或安全退出；游戏中可暂停、重开本关或返回菜单；清除敌人后进入下一关，生命归零进入失败界面，完成第三关显示最终统计。所有可见菜单按钮同时支持键盘和鼠标悬停/点击。",
        "玩法需求：玩家以四方向在固定网格中移动，脚底碰撞盒不能穿越固定墙、木箱、边界或已经离开的水泡格。水泡在2.5秒后产生十字水浪；固定墙完全阻挡，木箱被破坏并终止该方向传播，水浪可立即连锁触发其他水泡。",
        "敌人与关卡需求：第一关以巡逻敌人为主，第二关加入根据曼哈顿距离追踪玩家的敌人，第三关加入10点生命Boss。普通敌人与Boss共享Enemy接口，Boss每次有效水浪只损失一点生命，血量在顶部状态条和对象头顶同步显示。",
        "成长与反馈需求：木箱有概率掉落范围、容量、速度、护盾道具，强化值存在上限。HUD显示角色头像、生命、强化状态、分数、敌人/Boss血量、累计木箱、道具和用时。图片与音频缺失时程序不应崩溃，构建过程自动复制assets资源。",
    ]
    set_body(paragraph_by_text(document, "此部分对系统所有功能进行详细描述"), requirements[0])
    h23 = paragraph_by_text(document, "2.3 系统设计")
    for text in requirements[1:]:
        add_body_before(h23, text)

    ui_intro = paragraph_by_text(document, "此部分绘制系统界面的原型图并说明，例如：")
    set_body(ui_intro, "界面固定为960×720像素，左侧主区域用于地图或内容卡片，右侧228像素HUD显示角色和状态。视觉采用奶油纸白、深灰蓝轮廓、水泡蓝、珊瑚粉、薄荷绿与蜂蜜黄，所有按钮保持圆角和明确点击区域。角色原画、透明行走精灵和道具图集均已接入实际界面。")
    caption22 = paragraph_by_text(document, "图2.2 界面原型图")
    set_body(caption22, "图2.2 游戏完整UI原型图", False)
    caption22.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption22, PROJECT_ROOT / "assets" / "ui" / "ui_overview_v1.png", 5.55)

    flow_intro = paragraph_by_text(document, "此部分绘制系统流程图并说明，例如：")
    set_body(flow_intro, "主循环以场景枚举为状态机：窗口消息和InputManager先形成输入状态，当前场景处理输入后更新对象，随后完成碰撞、失效对象清理和绘制。暂停及结算场景不更新游戏对象，从而保证状态一致。")
    caption23 = paragraph_by_text(document, "图2.3 系统流程图")
    set_body(caption23, "图2.3 游戏主流程图", False)
    caption23.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption23, flow_diagram, 4.6)

    class_intro = paragraph_by_text(document, "如果使用C++面向对象开发，则此部分可以绘制单个类图或类关系图，例如：")
    class_text = (
        "项目明确落实封装、继承和多态。GameMap把地图数组设为私有，只允许通过查询、碰撞和破箱方法维护；Player把生命、速度、容量、范围、护盾与无敌时间设为私有，并通过行为方法保持上限和状态一致。GameObject定义update、draw、bounds纯虚接口，Character继承后复用移动与碰撞，Player和Enemy继续继承Character，PatrolEnemy、HunterEnemy、BossEnemy再继承Enemy。主循环使用vector<unique_ptr<Enemy>>保存异构对象，通过Enemy基类虚函数调用真实派生实现，因此新增敌人类型不需要复制主循环。"
    )
    set_body(class_intro, class_text)
    caption24 = paragraph_by_text(document, "图2.4 类关系图")
    set_body(caption24, "图2.4 主要类关系与STL容器", False)
    caption24.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption24, class_diagram, 5.55)
    h24 = paragraph_by_text(document, "2.4 系统实现")
    add_stl_table(document, h24)
    add_body_before(h24, "动态容器清理采用擦除-移除模式：对象先将active标记为false，统一阶段再用remove_if把失效元素移动到尾部并erase。敌人元素的unique_ptr随擦除自动释放派生对象，避免裸指针所有权和内存泄漏。")

    implementation_intro = paragraph_by_text(document, "此部分为最终实现效果图（多张图），例如：")
    set_body(implementation_intro, "以下画面由当前Release x64程序的--capture-report模式自动绘制并保存，和提交工程使用同一套渲染函数与资源，不是概念占位图。")
    caption25 = paragraph_by_text(document, "图2.5 游戏第一关")
    set_body(caption25, "图2.5 游戏主菜单", False)
    caption25.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption25, REPORT_ASSETS / "game_main_menu.png", 5.55)
    caption26 = paragraph_by_text(document, "图2.6 游戏第二关")
    set_body(caption26, "图2.6 差异化角色选择", False)
    caption26.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_before(caption26, REPORT_ASSETS / "game_character_select.png", 5.55)
    h242 = paragraph_by_text(document, "2.4.2 核心代码")
    add_image_before(h242, REPORT_ASSETS / "game_level1_items.png", 5.55)
    add_caption_before(h242, "图2.7 第一关地图、四类道具和HUD")
    add_image_before(h242, REPORT_ASSETS / "game_level3_boss.png", 5.55)
    add_caption_before(h242, "图2.8 第三关Boss与双血条显示")
    add_image_before(h242, REPORT_ASSETS / "game_victory.png", 5.55)
    add_caption_before(h242, "图2.9 最终通关与累计统计")

    code_intro = paragraph_by_text(document, "此部分为核心代码，需要排版整齐（不要截图），在关键行添加注释！ 例如：")
    set_body(code_intro, "以下代码节选自实际工程，分别展示抽象接口、角色配置、爆炸传播、多态受击和STL对象清理。代码以文本排版，关键逻辑添加中文注释。")
    h25 = paragraph_by_text(document, "2.5 系统测试")
    parent = code_intro._p.getparent()
    start_index = parent.index(code_intro._p) + 1
    end_index = parent.index(h25._p)
    for element in list(parent)[start_index:end_index]:
        if element.tag == qn("w:p"):
            parent.remove(element)

    snippets = [
        ("（1）抽象接口与运行时多态", r'''
class GameObject {
public:
    virtual ~GameObject() = default;
    virtual void update(float deltaTime) = 0;
    virtual void draw() const = 0;
    virtual RectF bounds() const = 0;
};

class Enemy : public Character {
public:
    // 普通敌人和Boss使用同一受击接口
    virtual EnemyHitResult takeWaveHit();
    virtual int scoreValue() const { return 300; }
protected:
    virtual Vec2 chooseDirection(...) = 0;
};'''),
        ("（2）角色属性封装与集中配置", r'''
void Player::resetForNewGame(CharacterStyle style) {
    const CharacterProfile profile = characterProfile(style);
    // 只能通过Player行为修改私有状态，外部不能直接破坏上限
    lives_ = profile.lives;
    bubbleCapacity_ = profile.bubbleCapacity;
    blastRange_ = profile.blastRange;
    shieldCharges_ = profile.shieldCharges;
    baseSpeed_ = profile.moveSpeed;
    speed_ = baseSpeed_;
    activeBubbles_ = 0;
}'''),
        ("（3）十字水浪传播与连锁触发", r'''
for (const GridPos& direction : directions) {
    for (int distance = 1; distance <= range; ++distance) {
        GridPos cell{ origin.row + direction.row * distance,
                      origin.column + direction.column * distance };
        TileType tile = map_.tileAt(cell);
        if (tile == TileType::SolidWall) break; // 固定墙完全阻挡
        waves_.emplace_back(cell);
        for (WaterBubble& other : bubbles_)
            if (other.active() && other.cell() == cell) other.trigger();
        if (tile == TileType::Crate) {           // 木箱破坏后停止
            map_.destroyCrate(cell);
            createPowerUp(cell);
            break;
        }
    }
}'''),
        ("（4）基类指针调用派生类受击逻辑", r'''
for (const std::unique_ptr<Enemy>& enemy : enemies_) {
    if (!enemy->active() || !intersects(wave.bounds(), enemy->bounds()))
        continue;
    // 虚函数自动分派到普通敌人或Boss实现
    EnemyHitResult result = enemy->takeWaveHit();
    if (result == EnemyHitResult::Defeated) {
        score_ += enemy->scoreValue();
        ++statistics_.enemiesDefeated;
    }
}'''),
        ("（5）STL擦除-移除与智能指针生命周期", r'''
enemies_.erase(
    std::remove_if(enemies_.begin(), enemies_.end(),
        [](const std::unique_ptr<Enemy>& enemy) {
            return !enemy->active();
        }),
    enemies_.end());
// erase后unique_ptr自动析构真实派生敌人，不需要手工delete'''),
    ]
    for title, code in snippets:
        add_body_before(h25, title, False)
        add_code_before(h25, code)

    test_intro = paragraph_by_text(document, "此部分为系统测试用例表，写10个以上测试用例，重点针对错误值和边界值进行测试，例如：")
    set_body(test_intro, "测试覆盖构建配置、界面资源、输入兼容、地图边界、碰撞、连锁爆炸、强化上限、图片相对路径、角色差异和Boss多段受击。自动构建与报告截图能够直接证明的项目标记为通过，其余交互边界在最终提交前按表逐项人工复核。")
    table_caption = paragraph_by_text(document, "表2.1 系统测试用例")
    set_body(table_caption, "表2.2 系统测试用例", False)
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    test_table = next(table for table in document.tables
                      if len(table.columns) == 6 and len(table.rows) >= 2)
    fill_test_table(test_table)

    summary = (
        "本项目完成了一个可独立运行的EasyX单人网格闯关游戏，并在课程要求的基础上加入四名差异化角色、多血量Boss、鼠标菜单、原创透明道具图标、行走精灵、背景音乐、音效、粒子效果、三关统计与自动报告截图。工程采用Visual Studio解决方案管理，Debug/Release和Win32/x64四种配置均能编译，资源通过项目目标自动复制，不依赖本机绝对路径。整体功能规模、界面完整度和可演示性达到最初设计目标。"
        "开发中最关键的技术难点有三个。第一是移动碰撞与操作手感：可见精灵比通道宽时容易卡墙，因此将碰撞范围收缩为脚底区域，并在转角处按格中心进行有限纠偏；动画方面在左右跨步之间插入回中帧，降低摆动。第二是动态对象与连锁爆炸：水泡、水浪、道具、粒子和敌人的数量不断变化，通过vector统一存储、active延迟失效和remove_if集中清理，避免遍历中删除；水浪传播严格区分固定墙、木箱和其他水泡。第三是面向对象扩展：使用GameObject、Character、Enemy三级抽象复用公共能力，Boss通过重写受击、得分和AI接口加入主循环，不需要在每个更新阶段判断具体类型。"
        "项目对封装、继承和多态进行了实际应用，而不是只在报告中描述。地图、玩家与对象状态均为私有数据，公开方法负责保持边界和上限；派生敌人复用基类移动、碰撞与危险格判断；vector<unique_ptr<Enemy>>使普通敌人和Boss能够在同一容器中运行时分派。std::array适合固定地图和键盘状态，std::vector适合动态对象，std::mt19937使随机关卡可复现，标准算法减少了手写查找和清理代码。"
        "目前仍可继续改进：AI使用局部评分而不是完整寻路，在复杂木箱布局下可能选择次优路线；音乐使用简单波形生成，层次和循环衔接仍可提升；报告测试表中的鼠标、极限碰撞和Boss连续命中项目需要在提交前完成最后人工点测。后续可以加入A*寻路、关卡编辑器、可保存最高分、更多Boss阶段和资源预加载缓存。在本次实践中，我从基本图元、消息输入和双缓冲逐步过渡到完整游戏循环与对象体系，理解了图形程序不仅是把画面画出来，更重要的是让数据结构、对象职责、状态转换、资源路径和测试证据保持一致。"
    )
    set_body(paragraph_by_text(document, "总结项目总体完成情况、技术难点、解决方法以及不足等，不少于500字。"), summary)

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    document.core_properties.title = "面向对象方法程序设计与实践报告 - 萌泡大作战"
    document.core_properties.subject = "C++、EasyX、封装、继承、多态与STL综合实践"
    document.core_properties.author = args.student_name
    document.save(output)

    week1_count = count_han(week1_intro + week1_difficult + week1_cases)
    week2_count = count_han(week2_intro + week2_difficult + week2_cases)
    summary_count = count_han([summary])
    if week1_count < 1000 or week2_count < 1000 or summary_count < 500:
        raise RuntimeError(f"word-count gate failed: week1={week1_count}, week2={week2_count}, summary={summary_count}")
    print(f"WROTE {output}")
    print(f"HAN_COUNTS week1={week1_count} week2={week2_count} summary={summary_count}")
    return output


def parse_args():
    parser = argparse.ArgumentParser(description="Generate the course report from the retained teacher template.")
    parser.add_argument("--template", default=str(WORKSPACE_ROOT / "面向对象方法程序设计与实践-报告模版.docx"))
    parser.add_argument("--output", default=str(REPORT_DIR / "程序设计课程实践报告-待填写姓名.docx"))
    parser.add_argument("--sequence", default="[请填写序号]")
    parser.add_argument("--student-id", default="[请填写学号]")
    parser.add_argument("--student-name", default="[请填写姓名]")
    parser.add_argument("--teacher", default="[请填写任课教师]")
    parser.add_argument("--material-link", default="[待填写]")
    return parser.parse_args()


if __name__ == "__main__":
    build_report(parse_args())
