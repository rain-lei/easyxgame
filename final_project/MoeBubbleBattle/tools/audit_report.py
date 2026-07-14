#!/usr/bin/env python3
"""Structural QA for the generated course-practice report."""

from __future__ import annotations

import argparse
import hashlib
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


EXPECTED_TEMPLATE_SHA256 = "C0C07D2861D80EE92DD4A181811CA3374CF96C38EAD6F6D827DCA70C4B00607B"
REQUIRED_HEADINGS = [
    "1 技术调研报告",
    "1.1 第一周学习总结（除去代码，不少于1000汉字）",
    "1.1.1 内容简介",
    "1.1.2 难点和解决办法",
    "1.1.3 学习案例",
    "1.2 第二周学习总结（除去代码，不少于1000汉字）",
    "1.2.1 内容简介",
    "1.2.2 难点和解决办法",
    "1.2.3 学习案例",
    "2 项目开发报告",
    "2.1 项目简介",
    "2.2 需求分析",
    "2.3 系统设计",
    "2.4 系统实现",
    "2.5 系统测试",
    "2.6 项目总结",
]
FORBIDDEN_TEMPLATE_TEXT = [
    "本周所学内容",
    "学习中遇到的问题以及个人是如何解决的",
    "此部分对所选项目进行简要说明",
    "此部分要求绘制系统功能模块图",
    "此部分为最终实现效果图",
    "贪吃蛇",
    "推箱子",
]
REQUIRED_PACKAGE_PARTS = [
    "word/styles.xml",
    "word/settings.xml",
    "word/theme/theme1.xml",
    "word/header1.xml",
    "word/footer1.xml",
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def count_han(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def paragraph_index(document: Document, exact_text: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == exact_text:
            return index
    raise ValueError(exact_text)


def section_han_count(document: Document, start: str, end: str) -> int:
    start_index = paragraph_index(document, start)
    end_index = paragraph_index(document, end)
    return count_han("\n".join(p.text for p in document.paragraphs[start_index + 1:end_index]))


def parse_args() -> argparse.Namespace:
    project = Path(__file__).resolve().parents[1]
    workspace = project.parents[1]
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path,
                        default=workspace / "面向对象方法程序设计与实践-报告模版.docx")
    parser.add_argument("--report", type=Path,
                        default=project / "report" / "程序设计课程实践报告-待填写姓名.docx")
    parser.add_argument("--expect-filled", action="store_true",
                        help="require all cover/material placeholders to be filled")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    template = args.template.resolve()
    report = args.report.resolve()
    errors: list[str] = []

    if not template.exists() or not report.exists():
        print(f"ERROR missing file: template={template.exists()} report={report.exists()}")
        return 1

    template_hash = sha256(template)
    if template_hash != EXPECTED_TEMPLATE_SHA256:
        errors.append(f"teacher template hash changed: {template_hash}")

    document = Document(report)
    texts = [paragraph.text for paragraph in document.paragraphs]
    full_text = "\n".join(texts)

    missing_headings = [heading for heading in REQUIRED_HEADINGS if heading not in texts]
    if missing_headings:
        errors.append("missing headings: " + ", ".join(missing_headings))

    residual = [item for item in FORBIDDEN_TEMPLATE_TEXT if item in full_text]
    if residual:
        errors.append("template instructions/examples remain: " + ", ".join(residual))

    table_shapes = [(len(table.rows), len(table.columns)) for table in document.tables]
    expected_table_shapes = [(22, 5), (6, 3), (21, 6)]
    if table_shapes != expected_table_shapes:
        errors.append(f"unexpected tables: {table_shapes}")

    if len(document.inline_shapes) != 13:
        errors.append(f"expected 13 report images, got {len(document.inline_shapes)}")

    section = document.sections[0]
    page_inches = (round(section.page_width.inches, 2), round(section.page_height.inches, 2))
    margins = tuple(round(value.inches, 2) for value in (
        section.left_margin, section.right_margin, section.top_margin, section.bottom_margin
    ))
    if page_inches != (8.27, 11.69):
        errors.append(f"page size changed: {page_inches}")
    if margins != (1.25, 1.25, 1.0, 1.0):
        errors.append(f"margins changed: {margins}")

    week1_count = section_han_count(
        document,
        "1.1 第一周学习总结（除去代码，不少于1000汉字）",
        "1.2 第二周学习总结（除去代码，不少于1000汉字）",
    )
    week2_count = section_han_count(
        document, "1.2 第二周学习总结（除去代码，不少于1000汉字）", "2 项目开发报告"
    )
    summary_count = count_han("\n".join(
        texts[paragraph_index(document, "2.6 项目总结") + 1:]
    ))
    if week1_count < 1000:
        errors.append(f"week 1 below 1000 Han characters: {week1_count}")
    if week2_count < 1000:
        errors.append(f"week 2 below 1000 Han characters: {week2_count}")
    if summary_count < 500:
        errors.append(f"summary below 500 Han characters: {summary_count}")

    with zipfile.ZipFile(report) as archive:
        names = set(archive.namelist())
        missing_parts = [part for part in REQUIRED_PACKAGE_PARTS if part not in names]
        if missing_parts:
            errors.append("missing preserved package parts: " + ", ".join(missing_parts))
        document_xml = archive.read("word/document.xml")
        settings_xml = archive.read("word/settings.xml")
        media_count = len([name for name in names if name.startswith("word/media/")])
        if media_count != 13:
            errors.append(f"expected 13 packaged media files, got {media_count}")

    root = etree.fromstring(document_xml)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    field_text = " ".join(root.xpath("//w:instrText/text()", namespaces=namespace))
    pageref_count = field_text.count("PAGEREF")
    if pageref_count != 24:
        errors.append(f"expected 24 TOC PAGEREF fields, got {pageref_count}")
    if b"updateFields" not in settings_xml:
        errors.append("w:updateFields is missing")

    placeholders = sorted(set(re.findall(r"\[[^\]]+\]", full_text)))
    expected_placeholders = [] if args.expect_filled else [
        "[待填写]", "[请填写任课教师]", "[请填写姓名]", "[请填写学号]", "[请填写序号]"
    ]
    if placeholders != expected_placeholders:
        errors.append(f"unexpected placeholders: {placeholders}")

    print(f"REPORT {report}")
    print(f"TEMPLATE_SHA256 {template_hash}")
    print(f"REPORT_SHA256 {sha256(report)}")
    print(f"PARAGRAPHS {len(document.paragraphs)}")
    print(f"TABLES {table_shapes}")
    print(f"IMAGES inline={len(document.inline_shapes)} media={media_count}")
    print(f"FIELDS pageref={pageref_count} updateFields={b'updateFields' in settings_xml}")
    print(f"PAGE inches={page_inches} margins={margins}")
    print(f"HAN_COUNTS week1={week1_count} week2={week2_count} summary={summary_count}")
    print(f"PLACEHOLDERS {placeholders}")

    if errors:
        print("AUDIT FAILED")
        for error in errors:
            print("- " + error)
        return 1
    print("AUDIT PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
